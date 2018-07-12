import sys
import urllib2
import os
import io
from docx import Document
from docx.shared import RGBColor
import time
from selenium import webdriver
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.firefox.firefox_binary import FirefoxBinary
import time
import datetime 
from termcolor import colored

# global var 
keyword_file_name = "keywords"
extension = ".txt"
job_key = "job-"
delimiter = "-"
roll_back_task = list() 

def roll_back(path , keyword_file , job_file):
    
    current_dir = os.getcwd()
    job_path =path+job_file
    key_path = path + keyword_file 
    os.system('mv' + " " +job_path+" "+current_dir)
    os.system('rm' + " " +key_path)
    if not os.listdir(path) :
        os.system('rmdir' + " " + path)
    
    
def match_rate_writer(before , after, path , position):
    import os.path
    completeName = os.path.join(path, "match_rate-" + position + " .txt")         
    text_file = open(completeName, "w")
    text_file.write("match rate before: " + before )
    text_file.write("match rate after: " + after )
    text_file.close()
    
def webdriverwait(driver):
    
    #element = WebDriverWait(driver,50).until(
    #lambda x: x.find_element_by_id("matchRateRadialScore"))
    time.sleep(20)
    element = driver.find_element_by_id("matchRateRadialScore") 
    element_attribute_value = element.get_attribute('data-progress')
    return element_attribute_value
    
def set_env(directory,position_name):
    path = os.getcwd()
     
    if not os.path.exists(directory):
        os.makedirs(directory)
    keyword_file = keyword_file_name+"-"+position_name+extension
    open(path+"/"+directory+"/"+keyword_file, 'w')
    path = path+"/"+directory
    return path+"/",keyword_file

def move_job_desc(path,file_name): 
    os.system('mv' + " " +file_name+" "+path)

def screen_shoot(path,pic_name,after=False):
    if after : 
        os.system('gnome-screenshot -f' + " " + path+pic_name + "_resultAfter.png")
    else: 
        os.system('gnome-screenshot -f' + " " + path+pic_name + "_resultBefore.png")

def read_resume_job(resume_path,job_path):

    # read job description
    file_job_description = io.open(job_path,  mode="r", encoding="utf-8")
    job_description = file_job_description.read()
    # read resume 
    resume = " " 
    document = Document(resume_path)
    for para in document.paragraphs:
        resume = resume + " " + para.text
    return resume , job_description

def selinume_request(resume,job_description):
  
    binary = FirefoxBinary('/usr/bin/firefox')
    driver = webdriver.Firefox(firefox_binary=binary)
    driver.get('https://www.jobscan.co/')

    text_area_cv = driver.find_element_by_id('cv')
    text_area_cv.send_keys(resume) 

    text_area_jd = driver.find_element_by_id('jd')
    text_area_jd.send_keys(job_description)

    btnAnalyze = driver.find_element_by_id('btnAnalyze')
    btnAnalyze.click()

    return driver

def html_parser(site):
    
    print "parsing html"
    result = [] 
    #site= "https://www.jobscan.co/results/3580546"
    hdr = {'User-Agent': 'Mozilla/5.0'}
    req = urllib2.Request(site,headers=hdr)
    page = urllib2.urlopen(req)
    try: 
        from BeautifulSoup import BeautifulSoup
    except ImportError:
        from bs4 import BeautifulSoup
    parsed_html = BeautifulSoup(page , "lxml")
    tables = parsed_html.body.findAll('table')
    #finding_jobtitle_found = parsed_html.body.find("i", {"id": "finding-jobtitle-found"}).text
   # x = parsed_html.body.find("div", {"id": "matchRateRadialScore"})
    for table in tables : 
        tr = table.findAll("tr")
        for i in xrange(len(tr)): 
            if i > 0 :
                td = tr[i].findAll("td") 
                result.append(td[0].text + " " +  td[2].text + " "+ td [3].text)
    #result.append(finding_jobtitle_found + " " + "0" + " "+ "2")
    return result

def key_word_extraction(keywords_input,path,keyword_file):
    
    print "extracting keywords "
    #file_path = sys.argv[3]
    file_path = path + keyword_file
    array = []
    baseCount = 0
    for line in keywords_input:
        baseCount = baseCount +1
        result = line 
        result = result.strip()
        word_count = result[-1:]
        if word_count.isdigit() :
            word_count = int(word_count)
        else :
            word_count = 1 
        result = ''.join([i for i in line if not i.isdigit()])
        result = result.strip()
        result = result + ","
        for i in range(word_count):
            array.append(result)
    lines = ""
    count = 0 
    for key in array : 
        lines = lines + key
        count = count + 1
    lines = lines + "M.Sc. in Computer Science,"
    lines = lines + "Master degree in Computer Science"
    print str(baseCount) + " key words founded in file"
    print str(len(array)) + " key words appended to file located in " + file_path  
    print "dont forget to add position title" 
    with open(file_path,'w') as output_file:
        output_file.write(lines)
    return lines

def save_to_resume(str_keyword,resume,company_name,position_name,path):

    # paragraph number 46 is where the keyword appended to be spicific its under Personal section at the end of last line (High problem solving skills.)
    paragraph_number = 46
    #file_path  = sys.argv[1]
    #company_name = sys.argv[4]
    document = Document(resume)
    paragraphs = document.paragraphs
   # count = 0 
    #for paragraph in paragraphs:
     #   runs = paragraph.runs
      #  for run in runs:
       #     print run.text
        #print count
        #count = count + 1
    paragraphs[paragraph_number]._p.clear()

    #Recreate second paragraph
    run = paragraphs[paragraph_number].add_run(str_keyword)
    font = run.font
    from docx.shared import Pt
    font.name = 'Calibri'
    font.size = Pt(5)
    font.color.rgb = RGBColor(255,255,255)
    new_file_path = resume.replace(".docx","")
    new_file_path = path + new_file_path +"_"+company_name + "_" + position_name + ".docx"
    document.save(new_file_path)
    print "saved in " + new_file_path    
    return new_file_path

def convert_to_pdf(file_path):

    os.system('doc2pdf' + " " +file_path)
    print "saving pdf "

def run(resume_path,company_name,position_name,path,keyword_file,job_file):
    
    #resume_path = sys.argv[1]
    #job_path = sys.argv[2]
    #company_name = sys.argv[4]
    job_path =path+job_file 
    result = read_resume_job(resume_path,job_path)
    driver = selinume_request(result[0],result[1])
    match_rate_before = webdriverwait(driver)
    if int(match_rate_before) == 0 : 
        roll_back_task.append([path , keyword_file , job_file])
        print colored("task failed wait for rollback in the end " , 'red')
        return
    site = driver.current_url
    screen_shoot(path,company_name+"-"+position_name)
    keywords_input = html_parser(site)
    driver.close()
    str_keyword = key_word_extraction(keywords_input,path,keyword_file)
    new_resume_path = save_to_resume(str_keyword,resume_path,company_name,position_name,path)
    print "sending resume again to check rate after appended the key word ! good luck "
    result = read_resume_job(new_resume_path,job_path)
    driver = selinume_request(result[0],result[1])
    match_rate_after = webdriverwait(driver)
    screen_shoot(path,company_name+"-"+position_name,True)
    driver.close()
    convert_to_pdf(new_resume_path)
    match_rate_writer(match_rate_before,match_rate_after,path,position_name)

if __name__ == "__main__":
    
    start = time.time()
    resume_path = sys.argv[1]
    count = 0;
    for file in os.listdir(os.getcwd()):
        if file.endswith(".txt") and file.find("job-") == 0:
            start_task_timer = time.time()
            job_file = file
            file = file.replace(job_key,"")
            file = file.replace(extension,"")
            file = file.split(delimiter)
            company_name = file[0]
            position_name = file[1]
            print "start task " + company_name + " " + position_name 
            env = set_env(company_name,position_name)
            path = env[0]
            keyword_file = env[1]
            move_job_desc(path,job_file)
            try:
              run(resume_path,company_name,position_name,path,keyword_file,job_file)
            except:
               roll_back_task.append([path , keyword_file , job_file])
               print colored("task failed wait for rollback in the end " , 'red')
            count = count + 1;
            end_task_timer = time.time()
            total_time_task = end_task_timer - start_task_timer
            print "task " + company_name + " " + position_name + " taked" + str(datetime.timedelta(seconds=total_time_task))
            print "!!! DONE !!!"
            print "---------------------------------------------------------------------------------------"
    end = time.time()
    total_time = (end - start)
    avg_time = total_time / count if count > 0 else 0;
    print str(count) + " tasks taked : " + str(datetime.timedelta(seconds=total_time)) + " avg time : " + str(datetime.timedelta(seconds=avg_time))
    
    if len(roll_back_task) > 0 : 
        print colored("Rolling back" , "yellow")

    for roll in roll_back_task : 
      
        roll_back(roll[0],roll[1],roll[2]) 
 
